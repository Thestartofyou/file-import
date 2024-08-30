import os
from docx import Document
from pdf2docx import Converter

def convert_docx_to_pdf(docx_file, output_pdf):
    from docx2pdf import convert
    convert(docx_file, output_pdf)
    print(f"Converted {docx_file} to {output_pdf}")

def convert_pdf_to_docx(pdf_file, output_docx):
    converter = Converter(pdf_file)
    converter.convert(output_docx)
    converter.close()
    print(f"Converted {pdf_file} to {output_docx}")

def convert_txt_to_docx(txt_file, output_docx):
    doc = Document()
    with open(txt_file, 'r') as file:
        doc.add_paragraph(file.read())
    doc.save(output_docx)
    print(f"Converted {txt_file} to {output_docx}")

def convert_docx_to_txt(docx_file, output_txt):
    doc = Document(docx_file)
    with open(output_txt, 'w') as file:
        for para in doc.paragraphs:
            file.write(para.text + '\n')
    print(f"Converted {docx_file} to {output_txt}")

def convert_document(input_file, output_file):
    input_ext = os.path.splitext(input_file)[1].lower()
    output_ext = os.path.splitext(output_file)[1].lower()

    if input_ext == '.docx' and output_ext == '.pdf':
        convert_docx_to_pdf(input_file, output_file)
    elif input_ext == '.pdf' and output_ext == '.docx':
        convert_pdf_to_docx(input_file, output_file)
    elif input_ext == '.txt' and output_ext == '.docx':
        convert_txt_to_docx(input_file, output_file)
    elif input_ext == '.docx' and output_ext == '.txt':
        convert_docx_to_txt(input_file, output_file)
    else:
        print("Conversion type not supported.")

if __name__ == "__main__":
    input_file = "input.docx"  # Change this to your input file
    output_file = "output.pdf"  # Change this to your desired output file
    convert_document(input_file, output_file)
